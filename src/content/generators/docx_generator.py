"""
Word document generator
"""

import logging
import os
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.content.generators.base_generator import BaseGenerator

class DocxGenerator(BaseGenerator):
    """Generator for Word documents (.docx)"""
    
    def generate(self, directory: str, filename: str, description: str,
                industry: str, language: str, role: Optional[str] = None,
                date_range_str: Optional[str] = None) -> bool:
        """
        Generate Word document content.
        
        Args:
            directory: Target directory for the file
            filename: Name of the file to create
            description: Description of the file content
            industry: Industry context
            language: Language to use
            role: Specific role within the industry (optional)
            date_range_str: Date range information (optional)
            
        Returns:
            True if file was successfully created, False otherwise
        """
        if not DOCX_AVAILABLE:
            logging.error("python-docx package is not available. Please install it with: pip install python-docx")
            return False
            
        file_path = self.get_file_path(directory, filename)
        
        # Create prompt for document content
        prompt = self.create_prompt(description, industry, language, role, "Word document", date_range_str)
        
        # Generate content using LLM
        content = self.llm_client.get_completion(prompt)
        
        if not content:
            logging.error(f"Failed to generate content for {filename}")
            return False
        
        try:
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading(description, level=1)
            # Add paragraphs
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Save document
            doc.save(file_path)
            return True
        except Exception as e:
            logging.error(f"Failed to create Word document {filename}: {e}")
            return False 