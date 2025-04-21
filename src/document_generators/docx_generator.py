#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import logging
import random
from datetime import date, datetime, timedelta
from typing import Dict, Tuple, Any
from pathlib import Path

# Import docx library
from docx import Document

# Import from local modules
from src.language_utils import get_date_format

def create_docx_document(file_path: str, file_name: str, description: str, industry: str, language: str = "en-US", role: str = None, file_examples: Dict[str, str] = None, date_range: Tuple[date, date] = None, model: str = None) -> bool:
    """Create a Word document with content appropriate for the file name and description"""
    try:
        document = Document()
        
        # Determine document type based on file name
        file_name_lower = file_name.lower()
        
        try:
            # Choose document type based on keywords in the filename
            if any(keyword in file_name_lower for keyword in ['report', 'analysis', 'assessment']):
                create_report_document(document, file_name, description, industry, role, date_range, language)
            elif any(keyword in file_name_lower for keyword in ['proposal', 'plan', 'strategy']):
                create_proposal_document(document, file_name, description, industry, role, date_range, language)
            elif any(keyword in file_name_lower for keyword in ['manual', 'guide', 'handbook', 'instruction']):
                create_manual_document(document, file_name, description, industry, role, date_range, language)
            elif any(keyword in file_name_lower for keyword in ['minutes', 'meeting', 'notes']):
                create_minutes_document(document, file_name, description, industry, role, date_range, language)
            elif any(keyword in file_name_lower for keyword in ['memo', 'memorandum', 'notice']):
                create_memo_document(document, file_name, description, industry, role, date_range, language)
            else:
                create_generic_document(document, file_name, description, industry, role, date_range, language)
        except Exception as e:
            logging.error(f"Error creating specific document type: {str(e)}")
            # Fallback to generic document if specific creation fails
            create_generic_document(document, file_name, description, industry, role, date_range, language)
        
        # Save the document
        document.save(os.path.join(file_path, file_name))
        logging.info(f"Successfully created Word document: {os.path.join(file_path, file_name)}")
        return True
    except Exception as e:
        logging.error(f"Error creating Word document: {str(e)}")
        return False

def create_report_document(document, file_name: str, description: str, industry: str, role: str = None, date_range: Tuple[date, date] = None, language: str = "en-US") -> None:
    """Create a report-style Word document"""
    # Add title
    document.add_heading(file_name, 0)
    
    # Add description as subtitle
    document.add_paragraph(description).style = 'Subtitle'
    
    # Add basic information section
    document.add_heading('Overview', level=1)
    document.add_paragraph(f'Industry: {industry}')
    
    if role:
        document.add_paragraph(f'Role: {role}')
    
    # Add date range if available
    if date_range:
        start_date_str = get_date_format(date_range[0], "prompt", language)
        end_date_str = get_date_format(date_range[1], "prompt", language)
        document.add_paragraph(f'Period: {start_date_str} - {end_date_str}')
    
    # Add report sections
    document.add_heading('Executive Summary', level=1)
    document.add_paragraph('This report provides a comprehensive analysis of the current situation and outlines recommendations for future actions.')
    
    document.add_heading('Findings', level=1)
    document.add_paragraph('The primary findings from our analysis include:')
    
    # Add bullets for sample findings
    for i in range(1, 4):
        p = document.add_paragraph('', style='List Bullet')
        p.add_run(f'Finding {i}: Sample finding related to {industry}')
    
    document.add_heading('Analysis', level=1)
    document.add_paragraph(f'Analysis of the {industry} industry reveals several key trends and opportunities:')
    
    # Add sample analysis paragraphs
    for i in range(1, 3):
        document.add_paragraph(f'Analysis paragraph {i}. This is sample content for the {industry} industry report.')
    
    document.add_heading('Recommendations', level=1)
    document.add_paragraph('Based on our findings, we recommend the following actions:')
    
    # Add numbered list for recommendations
    for i in range(1, 4):
        p = document.add_paragraph('', style='List Number')
        p.add_run(f'Recommendation {i}: A strategic recommendation for the {industry} industry')
    
    document.add_heading('Conclusion', level=1)
    document.add_paragraph(f'In conclusion, this report has identified key aspects of the {industry} industry and provided actionable recommendations.')

def create_proposal_document(document, file_name: str, description: str, industry: str, role: str = None, date_range: Tuple[date, date] = None, language: str = "en-US") -> None:
    """Create a proposal-style Word document"""
    # Add title
    document.add_heading(file_name, 0)
    
    # Add description as subtitle
    document.add_paragraph(description).style = 'Subtitle'
    
    # Add basic information section
    document.add_heading('Overview', level=1)
    document.add_paragraph(f'Industry: {industry}')
    
    if role:
        document.add_paragraph(f'Role: {role}')
    
    # Add date range if available
    if date_range:
        start_date_str = get_date_format(date_range[0], "prompt", language)
        end_date_str = get_date_format(date_range[1], "prompt", language)
        document.add_paragraph(f'Period: {start_date_str} - {end_date_str}')
    
    # Add proposal sections
    document.add_heading('Introduction', level=1)
    document.add_paragraph(f'This proposal outlines a comprehensive plan for addressing the needs of the {industry} industry.')
    
    document.add_heading('Objectives', level=1)
    document.add_paragraph('The primary objectives of this proposal are:')
    
    # Add bullets for objectives
    for i in range(1, 4):
        p = document.add_paragraph('', style='List Bullet')
        p.add_run(f'Objective {i}: A key objective related to {industry}')
    
    document.add_heading('Proposed Solution', level=1)
    document.add_paragraph(f'Our proposed solution for the {industry} industry includes the following components:')
    
    # Add numbered list for solution components
    for i in range(1, 4):
        p = document.add_paragraph('', style='List Number')
        p.add_run(f'Component {i}: A key component of our solution')
    
    document.add_heading('Implementation Plan', level=1)
    document.add_paragraph('The implementation will follow these phases:')
    
    # Add table for implementation phases
    table = document.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Phase'
    header_cells[1].text = 'Timeline'
    header_cells[2].text = 'Deliverables'
    
    # Add data rows
    for i in range(1, 4):
        row_cells = table.rows[i].cells
        row_cells[0].text = f'Phase {i}'
        row_cells[1].text = f'Month {i}-{i+2}'
        row_cells[2].text = f'Deliverable {i}'
    
    document.add_heading('Budget', level=1)
    document.add_paragraph('The proposed budget for this project is as follows:')
    
    # Add table for budget
    budget_table = document.add_table(rows=4, cols=2)
    budget_table.style = 'Table Grid'
    
    # Add headers
    header_cells = budget_table.rows[0].cells
    header_cells[0].text = 'Item'
    header_cells[1].text = 'Cost'
    
    # Add budget items
    total = 0
    for i in range(1, 3):
        cost = random.randint(10000, 50000)
        total += cost
        row_cells = budget_table.rows[i].cells
        row_cells[0].text = f'Budget Item {i}'
        row_cells[1].text = f'${cost:,}'
    
    # Add total
    row_cells = budget_table.rows[3].cells
    row_cells[0].text = 'Total'
    row_cells[1].text = f'${total:,}'
    
    document.add_heading('Conclusion', level=1)
    document.add_paragraph(f'This proposal presents a comprehensive solution for addressing the needs of the {industry} industry.')

def create_manual_document(document, file_name: str, description: str, industry: str, role: str = None, date_range: Tuple[date, date] = None, language: str = "en-US") -> None:
    """Create a manual/guide-style Word document"""
    # Add title
    document.add_heading(file_name, 0)
    
    # Add description as subtitle
    document.add_paragraph(description).style = 'Subtitle'
    
    # Add basic information section
    document.add_heading('Overview', level=1)
    document.add_paragraph(f'Industry: {industry}')
    
    if role:
        document.add_paragraph(f'Role: {role}')
    
    # Add date range if available
    if date_range:
        start_date_str = get_date_format(date_range[0], "prompt", language)
        end_date_str = get_date_format(date_range[1], "prompt", language)
        document.add_paragraph(f'Valid Period: {start_date_str} - {end_date_str}')
    
    # Add manual sections
    document.add_heading('Introduction', level=1)
    document.add_paragraph(f'This manual provides comprehensive guidelines and instructions for {industry} operations.')
    
    document.add_heading('Getting Started', level=1)
    document.add_paragraph('Follow these steps to get started:')
    
    # Add numbered list for steps
    for i in range(1, 4):
        p = document.add_paragraph('', style='List Number')
        p.add_run(f'Step {i}: Description of step {i}')
    
    # Add several instruction sections
    for i in range(1, 4):
        document.add_heading(f'Section {i}: Key Procedures', level=1)
        document.add_paragraph(f'This section covers important procedures for {industry}.')
        
        document.add_heading(f'Procedure {i}.1', level=2)
        document.add_paragraph('Follow these steps:')
        
        # Add steps for each procedure
        for j in range(1, 4):
            p = document.add_paragraph('', style='List Number')
            p.add_run(f'Step {j}: Description of step {j} for procedure {i}.1')
    
    document.add_heading('Troubleshooting', level=1)
    document.add_paragraph('Common issues and their solutions:')
    
    # Add table for troubleshooting
    trouble_table = document.add_table(rows=4, cols=2)
    trouble_table.style = 'Table Grid'
    
    # Add headers
    header_cells = trouble_table.rows[0].cells
    header_cells[0].text = 'Issue'
    header_cells[1].text = 'Solution'
    
    # Add issues and solutions
    for i in range(1, 4):
        row_cells = trouble_table.rows[i].cells
        row_cells[0].text = f'Common Issue {i}'
        row_cells[1].text = f'Solution for Issue {i}'
    
    document.add_heading('Glossary', level=1)
    document.add_paragraph('Key terms used in this manual:')
    
    # Add table for glossary
    glossary_table = document.add_table(rows=4, cols=2)
    glossary_table.style = 'Table Grid'
    
    # Add headers
    header_cells = glossary_table.rows[0].cells
    header_cells[0].text = 'Term'
    header_cells[1].text = 'Definition'
    
    # Add terms and definitions
    for i in range(1, 4):
        row_cells = glossary_table.rows[i].cells
        row_cells[0].text = f'Term {i}'
        row_cells[1].text = f'Definition of Term {i} in the context of {industry}'

def create_minutes_document(document, file_name: str, description: str, industry: str, role: str = None, date_range: Tuple[date, date] = None, language: str = "en-US") -> None:
    """Create a meeting minutes Word document"""
    # Add title
    document.add_heading(file_name, 0)
    
    # Add description as subtitle
    document.add_paragraph(description).style = 'Subtitle'
    
    # Get meeting date - use first date from date_range if available, or today
    meeting_date = date_range[0] if date_range else datetime.now().date()
    meeting_date_str = get_date_format(meeting_date, "prompt", language)
    
    # Add meeting details section
    document.add_heading('Meeting Details', level=1)
    document.add_paragraph(f'Date: {meeting_date_str}')
    document.add_paragraph(f'Industry: {industry}')
    
    if role:
        document.add_paragraph(f'Meeting Chair: {role}')
    
    document.add_paragraph('Location: Conference Room A')
    document.add_paragraph('Time: 10:00 AM - 11:30 AM')
    
    # Add attendees
    document.add_heading('Attendees', level=1)
    attendees = ['John Smith', 'Emma Johnson', 'Michael Lee', 'Sarah Chen', 'David Rodriguez']
    
    for attendee in attendees:
        p = document.add_paragraph('', style='List Bullet')
        p.add_run(attendee)
    
    document.add_heading('Agenda', level=1)
    agenda_items = ['Review of previous meeting minutes', 'Project updates', 'Budget discussion', 'Upcoming deadlines', 'Any other business']
    
    for item in agenda_items:
        p = document.add_paragraph('', style='List Number')
        p.add_run(item)
    
    document.add_heading('Discussion', level=1)
    
    # Add sections for each agenda item
    for i, item in enumerate(agenda_items, 1):
        document.add_heading(f'{i}. {item}', level=2)
        document.add_paragraph(f'Discussion about {item.lower()} in the context of the {industry} industry.')
        document.add_paragraph('Key points:')
        
        # Add key points for each agenda item
        for j in range(1, 3):
            p = document.add_paragraph('', style='List Bullet')
            p.add_run(f'Point {j}: Sample discussion point about {item.lower()}')
    
    document.add_heading('Action Items', level=1)
    document.add_paragraph('The following action items were identified:')
    
    # Add table for action items
    action_table = document.add_table(rows=4, cols=4)
    action_table.style = 'Table Grid'
    
    # Add headers
    header_cells = action_table.rows[0].cells
    header_cells[0].text = 'Action Item'
    header_cells[1].text = 'Assigned To'
    header_cells[2].text = 'Due Date'
    header_cells[3].text = 'Status'
    
    # Add action items
    for i in range(1, 4):
        due_date = meeting_date + timedelta(days=random.randint(7, 21))
        due_date_str = get_date_format(due_date, "prompt", language)
        
        row_cells = action_table.rows[i].cells
        row_cells[0].text = f'Action Item {i}'
        row_cells[1].text = random.choice(attendees)
        row_cells[2].text = due_date_str
        row_cells[3].text = 'Pending'
    
    document.add_heading('Next Meeting', level=1)
    next_meeting_date = meeting_date + timedelta(days=14)
    next_meeting_date_str = get_date_format(next_meeting_date, "prompt", language)
    document.add_paragraph(f'Date: {next_meeting_date_str}')
    document.add_paragraph('Time: 10:00 AM - 11:30 AM')
    document.add_paragraph('Location: Conference Room A')

def create_memo_document(document, file_name: str, description: str, industry: str, role: str = None, date_range: Tuple[date, date] = None, language: str = "en-US") -> None:
    """Create a memo/memorandum Word document"""
    # Add title
    document.add_heading('MEMORANDUM', 0)
    
    # Create memo header table
    table = document.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Add memo header details
    header_cells = table.rows[0].cells
    header_cells[0].text = 'TO:'
    header_cells[1].text = 'All Staff'
    
    header_cells = table.rows[1].cells
    header_cells[0].text = 'FROM:'
    header_cells[1].text = role if role else f'{industry} Management'
    
    header_cells = table.rows[2].cells
    header_cells[0].text = 'DATE:'
    memo_date = date_range[0] if date_range else datetime.now().date()
    memo_date_str = get_date_format(memo_date, "prompt", language)
    header_cells[1].text = memo_date_str
    
    header_cells = table.rows[3].cells
    header_cells[0].text = 'SUBJECT:'
    header_cells[1].text = file_name.replace('.docx', '')
    
    # Add a line break
    document.add_paragraph('')
    
    # Add description
    if description:
        document.add_paragraph(description)
    
    # Add memo content
    document.add_heading('Purpose', level=1)
    document.add_paragraph(f'This memorandum provides important information regarding {industry} operations.')
    
    document.add_heading('Background', level=1)
    document.add_paragraph(f'The {industry} industry has been experiencing significant changes that require attention.')
    document.add_paragraph('Additional context:')
    
    # Add bullets for background context
    for i in range(1, 3):
        p = document.add_paragraph('', style='List Bullet')
        p.add_run(f'Context point {i}: Important background information')
    
    document.add_heading('Key Information', level=1)
    document.add_paragraph('Please note the following important information:')
    
    # Add numbered list for key information
    for i in range(1, 4):
        p = document.add_paragraph('', style='List Number')
        p.add_run(f'Information point {i}: Critical information for all staff')
    
    document.add_heading('Action Required', level=1)
    document.add_paragraph('The following actions are required in response to this memorandum:')
    
    # Add bullets for required actions
    for i in range(1, 3):
        p = document.add_paragraph('', style='List Bullet')
        p.add_run(f'Action {i}: Required response from staff')
    
    document.add_heading('Timeline', level=1)
    document.add_paragraph('Please adhere to the following timeline:')
    
    # Add table for timeline
    timeline_table = document.add_table(rows=4, cols=2)
    timeline_table.style = 'Table Grid'
    
    # Add headers
    header_cells = timeline_table.rows[0].cells
    header_cells[0].text = 'Date'
    header_cells[1].text = 'Milestone'
    
    # Add timeline entries
    for i in range(1, 4):
        entry_date = memo_date + timedelta(days=i*7)
        entry_date_str = get_date_format(entry_date, "prompt", language)
        
        row_cells = timeline_table.rows[i].cells
        row_cells[0].text = entry_date_str
        row_cells[1].text = f'Milestone {i}'
    
    document.add_heading('Contact Information', level=1)
    document.add_paragraph('For questions or clarification, please contact:')
    document.add_paragraph('Name: John Smith')
    document.add_paragraph('Email: john.smith@example.com')
    document.add_paragraph('Phone: (555) 123-4567')

def create_generic_document(document, file_name: str, description: str, industry: str, role: str = None, date_range: Tuple[date, date] = None, language: str = "en-US") -> None:
    """Create a generic Word document with standard content"""
    # Add title
    document.add_heading(file_name, 0)
    
    # Add description as subtitle
    document.add_paragraph(description).style = 'Subtitle'
    
    # Add basic information section
    document.add_heading('Overview', level=1)
    document.add_paragraph(f'Industry: {industry}')
    
    if role:
        document.add_paragraph(f'Role: {role}')
    
    # Add date range if available
    if date_range:
        start_date_str = get_date_format(date_range[0], "prompt", language)
        end_date_str = get_date_format(date_range[1], "prompt", language)
        document.add_paragraph(f'Period: {start_date_str} - {end_date_str}')
    
    # Add generic sections
    document.add_heading('Introduction', level=1)
    document.add_paragraph(f'This document provides information related to the {industry} industry.')
    
    document.add_heading('Section 1', level=1)
    document.add_paragraph('This is the first main section of the document.')
    
    # Add subsections
    document.add_heading('Subsection 1.1', level=2)
    document.add_paragraph('Details for subsection 1.1:')
    
    # Add bullets
    for i in range(1, 4):
        p = document.add_paragraph('', style='List Bullet')
        p.add_run(f'Point {i}: Important information related to this subsection')
    
    document.add_heading('Subsection 1.2', level=2)
    document.add_paragraph('Details for subsection 1.2:')
    
    # Add numbered list
    for i in range(1, 4):
        p = document.add_paragraph('', style='List Number')
        p.add_run(f'Step {i}: Sequential step in a process')
    
    document.add_heading('Section 2', level=1)
    document.add_paragraph('This is the second main section of the document.')
    
    # Add a table
    document.add_paragraph('The following table summarizes key information:')
    table = document.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Category'
    header_cells[1].text = 'Description'
    header_cells[2].text = 'Value'
    
    # Add table content
    for i in range(1, 4):
        row_cells = table.rows[i].cells
        row_cells[0].text = f'Category {i}'
        row_cells[1].text = f'Description of category {i}'
        row_cells[2].text = f'Value {i}'
    
    document.add_heading('Conclusion', level=1)
    document.add_paragraph(f'This document has provided an overview of various aspects related to the {industry} industry.')
    document.add_paragraph('For more information, please contact the appropriate department.') 